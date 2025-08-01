import streamlit as st
from docx import Document
from docx.shared import Pt
import speech_recognition as sr
from pydub import AudioSegment
from io import BytesIO

st.title("Converter - Speech to Document")

audio_file = st.file_uploader("Upload your Hindi audio file (wav/mp3)", type=["wav", "mp3","m4a"])

if audio_file is not None:
    st.audio(audio_file)

    audio_bytes = BytesIO(audio_file.read())
    audio = AudioSegment.from_file(audio_bytes)
    wav_io = BytesIO()
    audio.export(wav_io, format="wav")
    wav_io.seek(0)

    recognizer = sr.Recognizer()
    with sr.AudioFile(wav_io) as source:
        audio_data = recognizer.record(source)
        try:
            transcription = recognizer.recognize_google(audio_data, language="hi-IN")
            edited_text = st.text_area("Edit the transcribed text here:", transcription, height=200)

            if st.button("Download as DOCX"):
                doc = Document()
                p = doc.add_paragraph()
                run = p.add_run(edited_text)
                run.font.size = Pt(16)  # Set font size; font name not set (Mangal font not on server)

                file_stream = BytesIO()
                doc.save(file_stream)
                file_stream.seek(0)
                st.download_button(
                    label="Download DOCX",
                    data=file_stream,
                    file_name="hindi_dictation.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

            st.info(
                "Note: The Mangal font is not embedded in the file since the server runs Linux. "
                "After downloading, open the document in MS Word and set the font to Mangal manually."
            )
        except Exception as e:
            st.error(f"Error during transcription: {e}")
else:
    st.info("Please upload an audio file to begin transcription.")
